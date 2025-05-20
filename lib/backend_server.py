# backend_server.py

import os
import uuid # To generate unique filenames for output
import docx # Make sure 'pip install python-docx'
from flask import Flask, request, jsonify, send_from_directory

# --- Configuration ---
UPLOAD_FOLDER = 'uploads'
OUTPUT_FOLDER = 'outputs' # Folder to store processed files
if not os.path.exists(UPLOAD_FOLDER):
    os.makedirs(UPLOAD_FOLDER)
if not os.path.exists(OUTPUT_FOLDER):
    os.makedirs(OUTPUT_FOLDER)

app = Flask(__name__)
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER
app.config['OUTPUT_FOLDER'] = OUTPUT_FOLDER
app.config['MAX_CONTENT_LENGTH'] = 16 * 1024 * 1024 # Limit file size (e.g., 16MB)

# Define the font substitutions (same as before)
FONT_SUBSTITUTIONS = {
    "Times New Roman": "EcoTimesNewRoman",
    "Arial": "EcoArial",
    "Calibri": "EcoCalibri"
}

# --- Font Substitution Logic (from previous example) ---
def substitute_fonts_in_docx(input_path, output_path, substitutions):
    """
    Reads a .docx file, substitutes specified fonts in text runs,
    and saves the result to a new file. Returns number of changes made.
    """
    try:
        document = docx.Document(input_path)
        fonts_changed_count = 0

        # Iterate through paragraphs and runs (main body)
        for paragraph in document.paragraphs:
            for run in paragraph.runs:
                original_font_name = run.font.name
                if original_font_name and original_font_name in substitutions:
                    new_font_name = substitutions[original_font_name]
                    run.font.name = new_font_name
                    run.font.element.rPr.rFonts.set(docx.oxml.ns.qn('w:eastAsia'), None) # Clear East Asia font if set
                    fonts_changed_count += 1

        # Iterate through Tables
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                         for run in paragraph.runs:
                            original_font_name = run.font.name
                            if original_font_name and original_font_name in substitutions:
                                new_font_name = substitutions[original_font_name]
                                run.font.name = new_font_name
                                run.font.element.rPr.rFonts.set(docx.oxml.ns.qn('w:eastAsia'), None)
                                fonts_changed_count += 1

        document.save(output_path)
        print(f"Substituted {fonts_changed_count} font instances.")
        return fonts_changed_count
    except Exception as e:
        print(f"Error during font substitution: {e}")
        raise # Re-raise the exception to be caught by the calling function


# --- API Endpoints ---

@app.route('/api/upload', methods=['POST'])
def upload_and_process_file():
    """
    API endpoint to receive a file upload, process it (DOCX only for now),
    and return results including savings and download link.
    """
    if 'file' not in request.files:
        return jsonify({"error": "No file part"}), 400
    file = request.files['file']
    if file.filename == '':
        return jsonify({"error": "No selected file"}), 400

    # --- Basic File Validation (Improve in production) ---
    if file and file.filename.lower().endswith('.docx'):
        try:
            # 1. Save uploaded file with a unique name (prevent collisions)
            original_filename = file.filename # Use secure_filename in production
            unique_id = str(uuid.uuid4())
            input_filename = f"{unique_id}_{original_filename}"
            input_path = os.path.join(app.config['UPLOAD_FOLDER'], input_filename)
            file.save(input_path)
            print(f"File saved temporarily to: {input_path}")

            # 2. Define output path
            output_filename = f"optimized_{input_filename}"
            output_path = os.path.join(app.config['OUTPUT_FOLDER'], output_filename)

            # 3. Process the file (font substitution)
            print(f"Processing {input_path} -> {output_path}...")
            changes_made = substitute_fonts_in_docx(input_path, output_path, FONT_SUBSTITUTIONS)
            print("Processing complete.")

            # 4. Calculate Placeholder Savings
            #    (In real app, base this on actual analysis or font metrics)
            #    For now: If any change was made, assume 20% ink saving.
            #    Page savings require layout analysis - reporting 0 for now.
            ink_savings_percent = 20 if changes_made > 0 else 0
            pages_saved = 0 # Placeholder - requires page count logic

            # 5. Prepare response
            #    Provide the filename needed for the download route
            response_data = {
                "message": "File processed successfully!",
                "originalFilename": original_filename,
                "processedFilename": output_filename, # Filename for download
                "savings": {
                    "inkPercent": ink_savings_percent,
                    "pagesSaved": pages_saved
                }
            }
            return jsonify(response_data), 200

        except Exception as e:
            print(f"Error during processing: {e}")
            # Clean up saved files if error occurs mid-processing? Maybe not input.
            # If output exists, maybe remove it: if os.path.exists(output_path): os.remove(output_path)
            return jsonify({"error": f"Processing failed: {e}"}), 500
        finally:
            # Optional: Clean up the original uploaded file? Or keep for logging?
            # if os.path.exists(input_path):
            #    os.remove(input_path)
            #    print(f"Cleaned up input file: {input_path}")
            pass

    else:
        # Handle non-docx files or missing files
        return jsonify({"error": "Invalid file type. Only .docx is supported currently."}), 400


# --- Download Endpoint ---
@app.route('/api/download/<filename>', methods=['GET'])
def download_file(filename):
    """
    Serves processed files from the OUTPUT_FOLDER.
    """
    print(f"Download request received for: {filename}")
    try:
        # Ensure the path is secure (prevents accessing files outside OUTPUT_FOLDER)
        return send_from_directory(app.config['OUTPUT_FOLDER'], filename, as_attachment=True)
    except FileNotFoundError:
        print(f"Error: File not found for download: {filename}")
        return jsonify({"error": "File not found"}), 404
    except Exception as e:
        print(f"Error during download: {e}")
        return jsonify({"error": "Could not process download request"}), 500


# --- Basic Status Endpoint Placeholder (Not used in this simplified flow) ---
@app.route('/api/status/<job_id>', methods=['GET'])
def get_status(job_id):
    # This is not used now because processing happens directly in upload
    # In a real app with background tasks, this would check job status.
    return jsonify({"status": "not_implemented"}), 501

if __name__ == '__main__':
    print("Starting Flask server (v2 - with processing)...")
    app.run(host='0.0.0.0', port=5000, debug=True)